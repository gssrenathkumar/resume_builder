from dotenv import load_dotenv
import streamlit as st
import os
import google.generativeai as genai
import pdfplumber
import fitz  # PyMuPDF
from io import BytesIO
import tempfile
from docx import Document
from docx.shared import Inches
import re
import all_functions as func
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import cv2
from PIL import Image
import streamlit as st
from zipfile import ZipFile
from io import BytesIO
import os
import shutil
from docx2pdf import convert
from PIL import Image, ImageDraw
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


#----------------------------
# Function to get response from text-based model for name query
def get_education_details_overall(input_text):
    input_prompt = """
        You are a data extracter. You have to extract the data from the given text. 
        You will have to answer the questions based on the input text.
        
    """
    name_query = "Your objective is to scrutinize the provided input text, identifying and extracting mentions of academic degrees and diplomas. With each qualification you find, you must also capture and include the name of the institute that awarded it, as well as the year the degree or diploma was awarded. Focus exclusively on academic degrees (e.g., B.A., M.S., Ph.D.) and diplomas. Organize your findings in a structured list format, where each entry is formatted as follows: 'Degree/Diploma, Institute Name, Year of Study.' Ensure precision in detailing the titles of degrees/diplomas, the names of the institutions, and the correct corresponding years. Omit any educational details that do not directly correspond to this format or that do not include academic degrees or diplomas."
    query = input_prompt + name_query
    model = genai.GenerativeModel("gemini-1.0-pro")
    response = model.generate_content([input_text, query])
    return response.text

#------------------------

def remove_pdf_and_docx_files_in_script_directory():
    """
    Remove all PDF, DOCX, PNG, JPEG, and JPG files from the directory where this script is located.
    """
    directory = os.path.dirname(os.path.abspath(__file__))
    for filename in os.listdir(directory):
        if filename.endswith((".pdf", ".docx", ".png", ".jpeg", ".jpg")):
            file_path = os.path.join(directory, filename)
            try:
                os.remove(file_path)
            except Exception as e:
                pass



def convert_to_pdf_if_docx(file_path):
    """
    Converts a .docx file to .pdf if the file is a .docx, 
    otherwise returns the file_path unchanged.
    
    Args:
    - file_path (str): Path to the file.
    
    Returns:
    - str: Converted file path if it's a .docx, else the original file path.
    """
    # Check if the file is a .docx
    if file_path.endswith('.docx'):
        # Define the PDF path
        pdf_path = file_path.rsplit('.', 1)[0] + '.pdf'

        # Load the .docx document
        document = Document(file_path)

        # Set up the PDF canvas
        c = canvas.Canvas(pdf_path, pagesize=letter)
        width, height = letter  # Get width and height of the page

        # Simple text positioning for demonstration purposes
        y_position = height - 72  # Start 1 inch from the top
        for paragraph in document.paragraphs:
            c.drawString(72, y_position, paragraph.text)
            y_position -= 12  # Move down 12 points for the next line

            # Check if we're near the bottom of the page
            if y_position < 72:
                c.showPage()  # Start a new page
                y_position = height - 72  # Reset the y position

        c.save()

        # Return the PDF path
        return pdf_path
    else:
        # Return the original file path if it's not a .docx file
        return file_path

# Function to delete directories
def delete_directories(directories):
    for directory in directories:
        if os.path.exists(directory):
            shutil.rmtree(directory)


# Function to zip a folder and return the zipped content as a BytesIO object
def zip_folder_to_bytesio(folder_path):
    bytes_io = BytesIO()
    with ZipFile(bytes_io, 'w') as zip_obj:
        for root, dirs, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                # Write file to zip, preserving the relative path
                zip_obj.write(file_path, os.path.relpath(file_path, start=folder_path))
    bytes_io.seek(0)  # Rewind the BytesIO object for reading
    return bytes_io


def input_imagedata(uploaded_file):
    if uploaded_file is not None:
        # Convert the uploaded file to an image
        image = Image.open(uploaded_file)
        return image
    else:
        raise FileNotFoundError("No file uploaded")
        
def get_gemini_response_image(input, image, query=None):
    model = genai.GenerativeModel("gemini-pro-vision")
    if query is not None:
        response = model.generate_content([input, image, query])
    else:
        response = model.generate_content([input, image])
    return response.text


def replace_hyphens_with_bullet_points(doc_name):
    document = Document(doc_name)
    for para in document.paragraphs:
        if '>' in para.text:
            para.text = para.text.replace('>', '\u2022')
        if 'None' in para.text:
            para.text = para.text.replace('None',' ')
        if 'NA' in para.text:
            para.text = para.text.replace('NA',' ')
    return document
    
def replace_symbol_with_dash(doc_name):
    document = Document(doc_name)
    for para in document.paragraphs:
        if '&!' in para.text:
            para.text = para.text.replace('&!', '-')
        elif '&' in para.text:
            para.text = para.text.replace("&",'-')
    return document


# Function to extract text from PDF using PyMuPDF
def extract_text_from_pdf(uploaded_file):
    resume_text = ''
    with tempfile.NamedTemporaryFile(delete=False) as temp_file:
        temp_file.write(uploaded_file.read())
        temp_filename = temp_file.name
    
    with fitz.open(temp_filename) as pdf_file:
        for page_num in range(len(pdf_file)):
            page = pdf_file.load_page(page_num)
            resume_text += page.get_text()
    
    os.unlink(temp_filename)  # Delete the temporary file
    return resume_text


# Function to get response from text-based model for name query
def get_name_response(input_text):
    try:
        input_prompt = """
            You are a resume analyzer. You have to extract the data from the resume text. 
            You will have to answer the questions based on the input resume text.
            
        """
        name_query = "What is the person name in the resume"
        query = input_prompt + name_query
        model = genai.GenerativeModel("gemini-pro")
        response = model.generate_content([input_text, query])
        return response.text
    except InvalidArgument as e:
        st.error("API Key is invalid. Please pass a valid API key.")
        st.stop()


# Function to get response from text-based model for summary query
def get_summary_response(input_text):
    input_prompt = """
        You are a resume analyzer. You have to extract the data from the resume text. 
        You will have to answer the questions based on the input resume text.
        
    """
    summary_query = ".Format is one paragraph.Generate one paragraph Experience Summary of 8 lines give it as the resume owner is giving his experience in summary.Dont give details of working place. Exclude personal pronouns such as I, Iam, Myself,meet etc. Expain what I have done.In case if personal pronouns are missing, dont not add one. State only facts.I want it in one paragraph."
    query = input_prompt + summary_query
    model = genai.GenerativeModel("gemini-pro")
    response = model.generate_content([input_text, query])
    return response.text


# Function to get response from text-based model for certifications query
def get_certifications_response(input_text):
    input_prompt = """
        You are a resume analyzer. You have to extract the data from the resume text. 
        You will have to answer the questions based on the input resume text.
        
    """
    certifications_query = """
Get all the certifications in the resume.
The Format is 

> Certificate 1
> Certificate 2
> Certificate 3
> Certificate 4

If there are no certificates, then print None.
"""

    model = genai.GenerativeModel("gemini-pro")
    query = input_prompt + certifications_query
    response = model.generate_content([input_text, query])
    return response.text


# Function to get response from text-based model for degree details query
def get_degree_details_response(input_text):
    input_prompt = """
        You are a resume analyzer. You have to extract the data from the resume text. 
        You will have to answer the questions based on the input resume text.
        
    """
    degree_query = "Extract degree name alone from the given text input, as per the given input text order.Format is degree1?degree2?degree3."
    model = genai.GenerativeModel("gemini-pro")
    query = input_prompt + degree_query
    response = model.generate_content([input_text, query])
    return response.text


# Function to get response from text-based model for education details query
def get_education_details_response(input_text):
    input_prompt = """
        You are a resume analyzer. You have to extract the data from the resume text. 
        You will have to answer the questions based on the input resume text.
        
    """
    education_query = "Get institution name from the given text input,as per the given input text order.Duplication is allowed.Format is institute1?institute2?institute3"


    
    model = genai.GenerativeModel("gemini-pro")
    query = input_prompt + education_query
    response = model.generate_content([input_text, query])
    return response.text

# Function to get response from text-based model for education years in descending order query
def get_education_years_response(input_text):
    input_prompt = """
        You are a resume analyzer. You have to extract the data from the resume text. 
        You will have to answer the questions based on the input resume text.
        
    """
    years_query = "Get the year of passing for each degree.The passing out year for each degree.Format is year1?year2?year3."
    model = genai.GenerativeModel("gemini-pro")
    query = input_prompt + years_query
    response = model.generate_content([input_text, query])
    full_text = str(response.text)
    year_pattern = r'\b[12]\d{3}\b'
    years = re.findall(year_pattern, full_text)
    formatted_years = '?'.join(years)
    return formatted_years

def get_technical_skills_response2(input_text):
    input_prompt = """
        You are a resume analyzer. You have to extract the data from the resume text. 
        You will have to answer the questions based on the input resume text.
        
    """
    skills_query1 = """
Retrieve a list of all technical skills mentioned in the resume. Format the list in bullet points with the ">" symbol at the beginning of each word. Each word should be on a new line.
"""


    model = genai.GenerativeModel("gemini-pro")
    query1 = input_prompt + skills_query1
    response = model.generate_content([input_text, query1])
    return str(response.text)


# Function to extract text from DOCX
def extract_text_from_docx(uploaded_file):
    resume_text = ''
    docx_file = Document(uploaded_file)
    for paragraph in docx_file.paragraphs:
        resume_text += paragraph.text + '\n'
    return resume_text

def get_omichannel_data(input_text):
    input_prompt = """
        You are a resume analyzer. You have to extract the data from the resume text. 
        You will have to answer the questions based on the input resume text.
        
    """
    query_1 = """ Please provide a concise summary of overview of the projects done for this resume.Give 4 points about the person project done.
"""
    model = genai.GenerativeModel("gemini-pro")
    query_ov = input_prompt + query_1
    response = model.generate_content([input_text, query_ov])
    return response.text

# Function to fill in the Word document template
def fill_invitation(template_path, output_path, name, summary, certifications):
    data = {
        '[First Name]': name,
        '[Experience Summary]': summary,
        '[Certifications]': certifications
    }
    
    doc = Document(template_path)
    
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, value)

    doc.save(output_path)


def fill_invitation2(template_path, output_path, summary, skills, project_experience,certifications1,summar):
    data = {
        '[summary_2]': summary,
        '[skills_2]': skills,
        '[project_experience_1]': project_experience,
        '[certificates_2]': certifications1,
        '[summar]':summar
    }
    doc = Document(template_path)
    
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, value)

    doc.save(output_path)

def fill_invitation3(template_path, output_path, name, summary):
    data = {
        '[name_3]': name,
        '[summary_3]': summary,
    }
    
    doc = Document(template_path)
    
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, value)

    doc.save(output_path)


    

# Function to fill in the degree details inside the table
def fill_table_degree_details(template_path, output_path, degree_details):
    data = {
        '[Degree Details]': degree_details,
    }

    doc = Document(template_path)

    # Iterate through tables in the document
    for table in doc.tables:
        # Check if the table contains the [Degree Details] placeholder
        contains_degree_details = False
        for cell in table._cells:
            if '[Degree Details]' in cell.text:
                contains_degree_details = True
                break
        
        if contains_degree_details:
            # Iterate through rows in the table
            for i, row in enumerate(table.rows):
                # Check if there are multiple degrees separated by '?'
                if '?' in data['[Degree Details]']:
                    degrees = data['[Degree Details]'].split('?')
                    # Check if the current row index is within the range of degrees
                    if i < len(degrees):
                        row.cells[0].paragraphs[0].text = degrees[i]
                    else:
                        row.cells[0].paragraphs[0].text = ''  # Empty for rows beyond the number of degrees
                else:
                    # Replace placeholders in the first cell of each row
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in data.items():
                                if key in paragraph.text:
                                    paragraph.text = paragraph.text.replace(key, value)

    # Save the modified document
    doc.save(output_path)



# Function to fill in the institute names inside the table
def fill_table_institute_details(template_path, output_path, institute_details):
    data = {
        '[Institute Name]': institute_details,
    }

    doc = Document(template_path)

    # Iterate through tables in the document
    for table in doc.tables:
        # Check if the table contains the [Institute Name] placeholder
        contains_institute_details = False
        for cell in table._cells:
            if '[Institute Name]' in cell.text:
                contains_institute_details = True
                break
        
        if contains_institute_details:
            # Iterate through rows in the table
            for i, row in enumerate(table.rows):
                # Check if there are multiple institute names separated by '?'
                if '?' in data['[Institute Name]']:
                    institutes = data['[Institute Name]'].split('?')
                    # Check if the current row index is within the range of institutes
                    if i < len(institutes):
                        row.cells[1].paragraphs[0].text = institutes[i]
                    else:
                        row.cells[1].paragraphs[0].text = ''  # Empty for rows beyond the number of institutes
                else:
                    # Replace placeholders in the second cell of each row
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in data.items():
                                if key in paragraph.text:
                                    paragraph.text = paragraph.text.replace(key, value)

    # Save the modified document
    doc.save(output_path)



# Function to fill in the education years inside the table
def fill_table_education_years(template_path, output_path, education_years):
    data = {
        '[Education Years]': education_years,
    }

    doc = Document(template_path)

    # Iterate through tables in the document
    for table in doc.tables:
        # Check if the table contains the [Education Years] placeholder
        contains_education_years = False
        for cell in table._cells:
            if '[Education Years]' in cell.text:
                contains_education_years = True
                break
        
        if contains_education_years:
            # Iterate through rows in the table
            for i, row in enumerate(table.rows):
                # Check if there are multiple education years separated by '?'
                if '?' in data['[Education Years]']:
                    years = data['[Education Years]'].split('?')
                    # Check if the current row index is within the range of education years
                    if i < len(years):
                        row.cells[2].paragraphs[0].text = years[i]
                    else:
                        row.cells[2].paragraphs[0].text = ''  # Empty for rows beyond the number of education years
                else:
                    # Replace placeholders in the third cell of each row
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in data.items():
                                if key in paragraph.text:
                                    paragraph.text = paragraph.text.replace(key, value)

    # Save the modified document
    doc.save(output_path)



# Function to fill in the skill set inside the existing table
def fill_table_skill_set(template_path, output_path, skill_set):
    data = {
        '[Skill Set]': skill_set,
    }
   
    doc = Document(template_path)
    
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, value)

    doc.save(output_path)



def get_work_experience_response(input_text):
    
    input_prompt = """
        You are a data extracter. You have to extract the data from the given text. 
        You will have to answer the questions based on the input text.
        
    """
    experience_query = """
Generate project details in the following format for each project in the work or internship experience:
Remeber the format given is very important which should be generated as per the given below
?/Project #1?/ \n
?/Title:?/ Project Title
?/Role :?/ What was the role of mine.If there are no role name specified, then print None
?/Period:?/ Project Period.Format is year1 &! year2
?/Technologies used :?/ Technologies used in that specific project.Predict from that particular project details. \n
?/Role and Responsibilities:?/
> Responsibility 1
> Responsibility 2
> Responsibility 3
Make these as points as it straightly send to docx.

Repeat this format for all projects. Please ensure the content is sourced from the provided data without assumptions.
"""
    query = input_prompt + experience_query
    model = genai.GenerativeModel("gemini-pro")
    response = model.generate_content([input_text, query])
    return response.text



def get_work_experience_response2(input_text):
    experience_query = """
Generate project details in the following format for each project in the work or internship experience:
Remeber the format given is very important which should be generated as per the given below
?/Project #1?/ \n
?/Organization:?/ Organization Name
?/Title:?/ Project Title
?/Role :?/ What was the role of mine.If there are no role name specified, then print None
?/Period:?/ Project Period.Format is year1 &! year2
?/Technologies used :?/ technologies used in this project.If there are no skills specified, then print None \n
?/Role and Responsibilities:?/
> Responsibility 1
> Responsibility 2
> Responsibility 3
Make these as points as it straightly send to docx.

Repeat this format for all projects. Please ensure the content is sourced from the provided data without assumptions.
"""
    model = genai.GenerativeModel("gemini-pro")
    response = model.generate_content([input_text, experience_query])
    return response.text
    
def relevant_project_experience(input_text):
    experience_query = """
Generate project details in the following format for each project in the work or internship experience:
Remeber the format given is very important which should be generated as per the given below
?/Project #1 ?/

?/Role :?/ (Short role name) \n
?/Role and Responsibilities ?/:
> Responsibility 1
> Responsibility 2
> Responsibility 3
Make these as points as it straightly send to docx.

Repeat this format for all projects. Please ensure the content is sourced from the provided data without assumptions.
"""

    model = genai.GenerativeModel("gemini-pro")
    response = model.generate_content([input_text, experience_query])
    return response.text


# Function to replace placeholder with organization count in the Word document
def remove_asterisks(text):
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        if line.startswith('*'):
            line = line[1:]  # Remove the asterisk at the beginning of the line
        cleaned_lines.append(line)
    cleaned_text = '\n'.join(cleaned_lines)
    return cleaned_text



def replace_organization_count(template_path, output_path, organization_count):
    data = {
        '[template]': organization_count
    }
    
    doc = Document(template_path)
    
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, value)

    doc.save(output_path)

def replace_organization_count2(template_path, output_path, organization_count):
    data = {
        '[project_2]': organization_count
    }
    
    doc = Document(template_path)
    
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, value)

    doc.save(output_path)


def bold_text_in_docx(file_path):
    def make_bold(text):
        if '?/' in text:
            parts = text.split('?/')
            new_text = ''
            for i in range(len(parts)):
                if i % 2 == 1:
                    new_text += parts[i].strip()
                else:
                    new_text += parts[i]
            return new_text
        else:
            return text.strip()

    def process_paragraph(paragraph):
        for run in paragraph.runs:
            if '?/' in run.text:
                run.text = make_bold(run.text)
                run.bold = True

    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)
    doc.save(file_path)


def remove_characters_from_docx(file_path):
    def remove_special_characters(text):
        # Replace characters with a blank space
        special_characters = ['*']
        for char in special_characters:
            text = text.replace(char, ' ')  # Replacing with a space
        return text

    # Load the document
    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = remove_special_characters(run.text)

    # Save the modifications back to the original document
    doc.save(file_path)

def remove_characters_from_docx2(file_path):
    def remove_special_characters(text):
        # Replace characters with a blank space
        special_characters = ['*','-']
        for char in special_characters:
            text = text.replace(char, ' ')  # Replacing with a space
        return text

    # Load the document
    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = remove_special_characters(run.text)

    # Save the modifications back to the original document
    doc.save(file_path)

def delete_rows_with_any_empty_cells(doc_path):
    # Load the Word document
    doc = Document(doc_path)
    
    # Iterate through each table in the document
    for table in doc.tables:
        # Reverse iterate over the rows because when we remove a row, it affects the index of subsequent rows
        for row in reversed(table.rows):
            # Check if any cell in the row is empty (contains no text or only whitespace)
            if any(cell.text.strip() == '' for cell in row.cells):
                # Remove the row
                # This accesses internal elements; be cautious as internal API may change
                table._element.remove(row._tr)
                
    # Save the modified document
    doc.save(doc_path)



# -------------------------


def pdf_to_image(pdf_path):
    # Split the PDF path into directory, base name, and extension
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    output_path = os.path.join(os.path.dirname(pdf_path), base_name + ".jpg")
    
    # Open the PDF
    pdf_document = fitz.open(pdf_path)
    
    # Get the first page
    first_page = pdf_document[0]
    
    # Convert the page to an image
    image = first_page.get_pixmap()
    
    # Save the image
    image.save(output_path)
    return output_path


def save_document(uploaded_file):
    if uploaded_file is not None:
        with open(uploaded_file.name, 'wb') as f:
            f.write(uploaded_file.getvalue())
        return uploaded_file.name
    else:
        st.warning("Please upload a document.")


def delete_document(filename):
    if os.path.exists(filename):
        os.remove(filename)
        st.success(f"{filename} deleted successfully.")
    else:
        st.warning("File does not exist.")


def extract_and_save_passport_photo(image_path):
    padding = 30
    output_file = "passport_photo.png"
    face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
    # Read the image
    image = cv2.imread(image_path)
    
    # Convert the image to grayscale
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    
    # Detect faces in the image
    faces = face_cascade.detectMultiScale(gray, scaleFactor=1.1, minNeighbors=3)
    
    # Check if a face is detected
    if len(faces) > 0:
        # Extract the first detected face
        x, y, w, h = faces[0]
        
        # Calculate coordinates for the passport photo with padding
        x1 = max(x - padding, 0)
        y1 = max(y - padding, 0)
        x2 = min(x + w + padding, image.shape[1])
        y2 = min(y + h + padding, image.shape[0])
        
        # Extract the passport photo region
        passport_photo = image[y1:y2, x1:x2]
        
        # Save the passport photo as a PNG file
        cv2.imwrite(output_file, passport_photo)
        print("Passport photo extracted and saved successfully!")
    else:
        print("No face detected in the input image.")



import os
from docx import Document
from docx.shared import Inches

def replace_placeholder_with_image(doc_path, image_path, image_width_inches=1.5):
    """
    Search for a placeholder in the document and replace it with an image or a blank
    if the image path is not valid.

    :param doc_path: Path to the .docx document.
    :param image_path: Path to the image to insert.
    :param image_width_inches: Width of the image in inches.
    """
    placeholder = "[photo]"
    # Load the document
    doc = Document(doc_path)
    
    # Check if the image file exists
    image_exists = os.path.exists(image_path)
    
    # Iterate through each paragraph to find the placeholder
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Found the placeholder, now clear all runs in the paragraph
            for run in paragraph.runs:
                run.clear()

            if image_exists:
                # Add the image in a new run in the same paragraph if the image exists
                run = paragraph.add_run()
                run.add_picture(image_path, width=Inches(image_width_inches))
            else:
                # If the image does not exist, add a new run with an empty string to effectively remove the placeholder
                paragraph.add_run('')
            break  # Remove this if you want to replace all occurrences

    # Save the modified document
    doc.save(doc_path)




def bold_text_between_markers11(paragraph, marker='?/'):
    """
    Make text between specified markers bold within a single paragraph.
    """
    run_objs = []  # List to keep track of runs to be formatted as bold
    in_bold = False  # Flag to track whether we're between the bold markers

    for run in paragraph.runs:
        # Split the run text by the marker; if marker is not found, we'll get the whole text back in a list of size 1
        parts = run.text.split(marker)
        
        # Iterate over parts. Odd-indexed parts are between markers
        for i, part in enumerate(parts):
            if i % 2 == 1:
                # This part should be bold
                new_run = paragraph.add_run(part)
                new_run.bold = True
                run_objs.append(new_run)
            else:
                # This part remains as is
                new_run = paragraph.add_run(part)
                run_objs.append(new_run)

        # Clear the original run text to prevent duplication
        run.text = ''

    # Re-add the formatted text to the paragraph, now including bold parts
    for run in run_objs:
        run.font.size = Pt(12)  # Example of setting font size; adjust as needed

def process_docx11(file_path):
    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        bold_text_between_markers11(paragraph)
    doc.save(file_path)  # Save the modified document


def zip_folder(folder_path):
    # Create a temporary file to hold the zip
    temp_dir = tempfile.mkdtemp()
    base_name = os.path.join(temp_dir, 'download')
    # Create a zip archive of the folder
    shutil.make_archive(base_name, 'zip', folder_path)
    # Return the path to the zipped file
    return base_name + '.zip'

